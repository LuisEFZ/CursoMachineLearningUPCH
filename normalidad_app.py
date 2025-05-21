# === IMPORTAR LIBRERÍAS NECESARIAS ===
import streamlit as st
import pandas as pd
import scipy.stats as stats
import matplotlib.pyplot as plt
import io
from docx import Document
import docx.shared
from fpdf import FPDF
import os

# === CONFIGURACIÓN DE LA PÁGINA ===
st.set_page_config(page_title="Pruebas de Normalidad", layout="centered")

st.title("🧪 Pruebas de Normalidad")
st.write("Sube un archivo Excel o CSV con hasta 100 datos numéricos para analizar si siguen una distribución normal.")

# === CARGA DE ARCHIVO ===
archivo = st.file_uploader("Carga tu archivo (.xlsx o .csv)", type=["xlsx", "csv"])

if archivo:
    # Verificar si tiene encabezado
    tiene_encabezado = st.checkbox("¿Tu archivo tiene encabezado?", value=True)

    try:
        if archivo.name.endswith(".csv"):
            df = pd.read_csv(archivo, header=0 if tiene_encabezado else None)
        else:
            df = pd.read_excel(archivo, header=0 if tiene_encabezado else None)
    except Exception as e:
        st.error(f"Error al leer el archivo: {e}")
        st.stop()

    # Seleccionar la columna que se va a analizar
    columna = st.selectbox("Selecciona la columna de datos", df.columns)

    # Convertir datos a numéricos y eliminar valores nulos
    datos_raw = pd.to_numeric(df[columna], errors='coerce')
    datos = datos_raw.dropna().astype(float).values[:100]

    # Mostrar cuántos datos se van a analizar
    st.write(f"Se analizarán **{len(datos)} datos**.")
    if len(datos) < 3:
        st.error("⚠️ Se requieren al menos 3 datos válidos para realizar las pruebas.")
        st.stop()

    # Selección de nivel de significancia
    st.subheader("⚙️ Configuración del análisis")
    opciones_sig = [0.10, 0.05, 0.01]
    seleccion_sig = st.selectbox("Selecciona un nivel de significancia", opciones_sig, index=1)
    sig_personalizado = st.number_input("...o ingresa otro valor", min_value=0.001, max_value=0.2, step=0.001, value=seleccion_sig)
    alpha = round(sig_personalizado, 4)
    st.write(f"Nivel de significancia seleccionado: **{alpha}**")

    # === PRUEBAS DE NORMALIDAD ===
    try:
        # Prueba de Shapiro-Wilk
        shapiro_stat, shapiro_p = stats.shapiro(datos)
        # Prueba de Kolmogorov-Smirnov sobre datos estandarizados
        ks_stat, ks_p = stats.kstest(stats.zscore(datos), 'norm')
        # Prueba de Anderson-Darling
        ad_result = stats.anderson(datos, dist='norm')
    except Exception as e:
        st.error(f"Error al ejecutar las pruebas estadísticas: {e}")
        st.stop()

    # Mostrar resultados
    st.subheader("📊 Resultados de las pruebas de normalidad")
    texto_resultado = []

    # Interpretar prueba de Shapiro-Wilk
    st.write("**Shapiro-Wilk**")
    st.write(f"Estadístico: {shapiro_stat:.4f}, Valor-p: {shapiro_p:.4f}")
    interpretacion_sw = "No se rechaza la normalidad ✅" if shapiro_p > alpha else "Se rechaza la normalidad ❌"
    st.write("→ " + interpretacion_sw)
    texto_resultado.append(f"Shapiro-Wilk: estadístico = {shapiro_stat:.4f}, p = {shapiro_p:.4f} → {interpretacion_sw}")

    # Interpretar prueba de Kolmogorov-Smirnov
    st.write("**Kolmogorov-Smirnov** (datos estandarizados)")
    st.write(f"Estadístico: {ks_stat:.4f}, Valor-p: {ks_p:.4f}")
    interpretacion_ks = "No se rechaza la normalidad ✅" if ks_p > alpha else "Se rechaza la normalidad ❌"
    st.write("→ " + interpretacion_ks)
    texto_resultado.append(f"Kolmogorov-Smirnov: estadístico = {ks_stat:.4f}, p = {ks_p:.4f} → {interpretacion_ks}")

    # Interpretar prueba de Anderson-Darling
    st.write("**Anderson-Darling**")
    st.write(f"Estadístico: {ad_result.statistic:.4f}")
    for cv, sig in zip(ad_result.critical_values, ad_result.significance_level):
        st.write(f"Nivel de significancia {sig:.1f}% → Valor crítico: {cv:.4f}")
    idx_crit = min(range(len(ad_result.significance_level)), key=lambda i: abs(ad_result.significance_level[i]/100 - alpha))
    if ad_result.statistic < ad_result.critical_values[idx_crit]:
        interpretacion_ad = f"No se rechaza la normalidad al {alpha:.2%} ✅"
    else:
        interpretacion_ad = f"Se rechaza la normalidad al {alpha:.2%} ❌"
    st.write("→ " + interpretacion_ad)
    texto_resultado.append(f"Anderson-Darling: estadístico = {ad_result.statistic:.4f} → {interpretacion_ad}")

    # === FUNCIÓN PARA CREAR GRÁFICOS ===
    def crear_graficos(datos):
        figuras = {}

        # Crear gráfico Q-Q
        fig_qq, ax_qq = plt.subplots()
        stats.probplot(datos, dist="norm", plot=ax_qq)
        ax_qq.set_title("Gráfico Q-Q")
        figuras["qq"] = fig_qq

        # Crear histograma
        fig_hist, ax_hist = plt.subplots()
        ax_hist.hist(datos, bins='auto', color='skyblue', edgecolor='black')
        ax_hist.set_title("Histograma")
        figuras["histograma"] = fig_hist

        # Crear boxplot
        fig_box, ax_box = plt.subplots()
        ax_box.boxplot(datos, vert=False)
        ax_box.set_title("Boxplot")
        figuras["boxplot"] = fig_box

        return figuras

    # Mostrar gráficos
    figuras = crear_graficos(datos)
    st.subheader("📈 Gráficos")
    for nombre, fig in figuras.items():
        st.pyplot(fig)

    # === EXPORTAR A WORD Y PDF ===
    st.subheader("📤 Exportar informe")
    if st.button("📄 Exportar resultados a Word y PDF"):

        # Resumen académico
        resumen = (
            "Este informe presenta los resultados del análisis de normalidad aplicado a una muestra de datos "
            "cuantitativos continuos. Se utilizaron tres pruebas estadísticas complementarias: Shapiro-Wilk, "
            "Kolmogorov-Smirnov y Anderson-Darling. Además, se incorporaron representaciones gráficas "
            "(histograma, boxplot y gráfico Q-Q) para una evaluación visual del ajuste a la distribución normal. "
            f"El nivel de significancia empleado fue de {alpha:.2%}, y los resultados permiten determinar, "
            "con base en criterios estadísticos, si la hipótesis de normalidad es razonable para la muestra analizada."
        )

        # Documento Word
        doc = Document()
        doc.add_heading("Informe de Pruebas de Normalidad", 0)
        doc.add_paragraph(resumen)
        doc.add_paragraph(f"\nArchivo analizado: {archivo.name}")
        doc.add_paragraph(f"Número de datos analizados: {len(datos)}\n")
        for linea in texto_resultado:
            doc.add_paragraph(linea)

        # Agregar gráficos al Word
        for nombre, fig in figuras.items():
            buf = io.BytesIO()
            fig.savefig(buf, format='png')
            buf.seek(0)
            doc.add_heading(nombre.capitalize(), level=1)
            doc.add_picture(buf, width=docx.shared.Inches(5))
            buf.close()

        buffer_word = io.BytesIO()
        doc.save(buffer_word)
        buffer_word.seek(0)

        # Documento PDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, "Informe de Pruebas de Normalidad", ln=True, align='C')
        pdf.ln(10)
        pdf.set_font("Arial", '', 12)
        for linea in resumen.split(". "):
            pdf.multi_cell(0, 10, linea.strip())
        pdf.ln(5)
        pdf.cell(0, 10, f"Archivo analizado: {archivo.name}", ln=True)
        pdf.cell(0, 10, f"Número de datos analizados: {len(datos)}", ln=True)
        pdf.ln(10)
        for linea in texto_resultado:
            pdf.multi_cell(0, 10, linea)
        pdf.ln(10)

        # Insertar gráficos al PDF
        for nombre, fig in figuras.items():
            buf = io.BytesIO()
            fig.savefig(buf, format='png')
            imagen_path = f"{nombre}.png"
            with open(imagen_path, "wb") as f:
                f.write(buf.read())
            pdf.add_page()
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(0, 10, nombre.capitalize(), ln=True, align='C')
            pdf.image(imagen_path, x=15, w=180)
            os.remove(imagen_path)

        buffer_pdf = io.BytesIO()
        pdf.output("temp_resultado.pdf")
        with open("temp_resultado.pdf", "rb") as f:
            buffer_pdf.write(f.read())
        buffer_pdf.seek(0)
        os.remove("temp_resultado.pdf")

        # Botones de descarga
        st.download_button("⬇️ Descargar informe Word", buffer_word, "resultado_normalidad.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.download_button("⬇️ Descargar informe PDF", buffer_pdf, "resultado_normalidad.pdf", mime="application/pdf")
